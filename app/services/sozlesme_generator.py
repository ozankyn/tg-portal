# -*- coding: utf-8 -*-
"""
Sözleşme DOCX Jeneratör (python-docx)

Admin .docx şablonu yükler → İK "Sözleşme Oluştur" dediğinde
{ad_soyad}, {ise_baslama} gibi placeholder'lar değiştirilir.
Format (font, bold, tablo) korunur.
"""
import io
import re
from datetime import date

from docx import Document


# ============================================================
# DEĞİŞKEN KATALOĞU
# ============================================================

OTOMATIK_DEGISKENLER = [
    {'kod': 'ad_soyad',          'aciklama': 'Çalışan ad soyad'},
    {'kod': 'tc_kimlik',         'aciklama': 'TC Kimlik No'},
    {'kod': 'ise_baslama',       'aciklama': 'İşe başlama tarihi (DD.MM.YYYY)'},
    {'kod': 'sozlesme_baslangic','aciklama': 'Sözleşme başlangıç tarihi'},
    {'kod': 'sozlesme_bitis',    'aciklama': 'Sözleşme bitiş tarihi'},
    {'kod': 'proje_adi',         'aciklama': 'Proje adı'},
    {'kod': 'pozisyon',          'aciklama': 'Pozisyon'},
    {'kod': 'departman',         'aciklama': 'Departman'},
    {'kod': 'adres',             'aciklama': 'Çalışan adresi'},
    {'kod': 'il',                'aciklama': 'İl'},
    {'kod': 'ilce',              'aciklama': 'İlçe'},
    {'kod': 'telefon',           'aciklama': 'Telefon'},
    {'kod': 'email',             'aciklama': 'E-posta'},
    {'kod': 'dogum_tarihi',      'aciklama': 'Doğum tarihi'},
    {'kod': 'dogum_yeri',        'aciklama': 'Doğum yeri'},
    {'kod': 'sicil_no',          'aciklama': 'Sicil numarası'},
    {'kod': 'bugunun_tarihi',    'aciklama': 'Bugünün (oluşturma) tarihi'},
]

MANUEL_DEGISKENLER = [
    {'kod': 'brut_ucret',     'aciklama': 'Brüt ücret (₺)',           'tip': 'number',  'default': ''},
    {'kod': 'net_ucret',      'aciklama': 'Net ücret (₺)',            'tip': 'number',  'default': ''},
    {'kod': 'calisma_gunu',   'aciklama': 'Haftalık çalışma günü',    'tip': 'number',  'default': '5'},
    {'kod': 'calisma_saati',  'aciklama': 'Günlük çalışma saati',     'tip': 'number',  'default': '9'},
    {'kod': 'deneme_suresi',  'aciklama': 'Deneme süresi (gün)',      'tip': 'number',  'default': '60'},
    {'kod': 'yillik_izin',    'aciklama': 'Yıllık izin günü',         'tip': 'number',  'default': '14'},
]


# ============================================================
# DEĞER ÇIKARIMI
# ============================================================

def _format_date(d):
    if not d:
        return ''
    if isinstance(d, str):
        return d
    return d.strftime('%d.%m.%Y')


def calisan_degiskenleri(calisan):
    """Çalışan modelinden otomatik değişken sözlüğünü çıkarır."""
    proje_adi = ''
    if calisan.kadro and calisan.kadro.proje:
        proje_adi = calisan.kadro.proje.ad

    return {
        'ad_soyad':           calisan.full_name or '',
        'tc_kimlik':          calisan.tc_kimlik or '',
        'ise_baslama':        _format_date(calisan.ise_baslama),
        'sozlesme_baslangic': _format_date(calisan.sozlesme_baslangic),
        'sozlesme_bitis':     _format_date(calisan.sozlesme_bitis),
        'proje_adi':          proje_adi,
        'pozisyon':           (calisan.pozisyon.ad if calisan.pozisyon else '') or '',
        'departman':          (calisan.departman.ad if calisan.departman else '') or '',
        'adres':              calisan.adres or '',
        'il':                 calisan.il or '',
        'ilce':               calisan.ilce or '',
        'telefon':            calisan.telefon or '',
        'email':              calisan.email or '',
        'dogum_tarihi':       _format_date(calisan.dogum_tarihi),
        'dogum_yeri':         calisan.dogum_yeri or '',
        'sicil_no':           calisan.sicil_no or '',
        'bugunun_tarihi':     _format_date(date.today()),
    }


# ============================================================
# PLACEHOLDER DOLDURMA
# ============================================================

_PLACEHOLDER_RE = re.compile(r'\{([a-zA-Z0-9_]+)\}')


def _replace_in_paragraph(paragraph, degerler):
    """Paragrafın run'ları arasına dağılmış {placeholder}'ları değiştirir.

    Word genelde placeholder'ı birden fazla run'a böler. Burada full text
    üzerinden eşleşmeleri buluyoruz; etkilenen run'ları sağdan sola
    güncelleyerek format'ı (ilk run'ın biçimi) koruyoruz.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = ''.join(r.text for r in runs)
    if '{' not in full_text:
        return

    matches = list(_PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    boundaries = []
    pos = 0
    for r in runs:
        boundaries.append([pos, pos + len(r.text), r])
        pos += len(r.text)

    for m in reversed(matches):
        start, end = m.span()
        value = str(degerler.get(m.group(1), ''))

        affected = [b for b in boundaries if b[0] < end and b[1] > start]
        if not affected:
            continue

        first = affected[0]
        last = affected[-1]
        fs, _fe, fr = first
        ls, _le, lr = last

        prefix = fr.text[:start - fs]
        suffix = lr.text[end - ls:]

        if first is last:
            fr.text = prefix + value + suffix
        else:
            fr.text = prefix + value
            for b in boundaries:
                if b is first or b is last:
                    continue
                if b[0] < end and b[1] > start:
                    b[2].text = ''
            lr.text = suffix

        pos = 0
        for b in boundaries:
            b[0] = pos
            b[1] = pos + len(b[2].text)
            pos = b[1]


def _replace_in_table(table, degerler):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, degerler)
            for nested in cell.tables:
                _replace_in_table(nested, degerler)


def _replace_in_doc_part(part, degerler):
    """Bir Document veya Header/Footer üzerindeki tüm paragraf + tabloları işler."""
    for paragraph in part.paragraphs:
        _replace_in_paragraph(paragraph, degerler)
    for table in part.tables:
        _replace_in_table(table, degerler)


def degiskenleri_doldur_docx(template_bytes, degerler):
    """.docx şablon bytes → doldurulmuş .docx bytes.

    Format korunur (bold, font, size, paragraf stili, tablo yapısı).
    Header/footer'lar da işlenir.
    """
    doc = Document(io.BytesIO(template_bytes))

    _replace_in_doc_part(doc, degerler)

    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is not None:
                _replace_in_doc_part(part, degerler)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def kullanilan_degiskenler_docx(template_bytes):
    """Şablonda kullanılan placeholder kodlarının unique listesi."""
    doc = Document(io.BytesIO(template_bytes))
    found = set()

    def _scan(part):
        for p in part.paragraphs:
            for kod in _PLACEHOLDER_RE.findall(p.text):
                found.add(kod)
        for t in part.tables:
            for row in t.rows:
                for cell in row.cells:
                    _scan(cell)

    _scan(doc)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is not None:
                _scan(part)

    return sorted(found)
